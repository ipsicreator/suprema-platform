from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\chris\Desktop\suprema-platform")
OUT = ROOT / "outputs" / "오류원인보고서.docx"


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_font(run, name: str = "Arial", size: int = 11, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=10)


def main() -> None:
    doc = Document()
    set_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("2028 입시계획 파일 오류 원인 보고서")
    set_font(r, size=18, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("대상 파일: 2028학년도_입시계획_최종_숫자만.xlsx")
    set_font(r, size=10)

    doc.add_paragraph("")

    h = doc.add_paragraph()
    r = h.add_run("1. 결론")
    set_font(r, size=12, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(
        "현재 파일은 최종본이 아니라, PDF/HWP/XLSX가 섞인 원문을 같은 규칙으로 처리하는 과정에서 "
        "열 매핑과 소스 식별이 흔들려 생긴 오염본입니다."
    )
    set_font(r)

    h = doc.add_paragraph()
    r = h.add_run("2. 오류가 난 직접 원인")
    set_font(r, size=12, bold=True)
    add_bullet(doc, "원문 식별 키(source_key)가 일부 파일에서 잘못 잡혀 학교별 소스 매칭이 틀어졌습니다.")
    add_bullet(doc, "PDF는 목차/본문/표가 섞여 추출되어, '모집인원'이 아닌 숫자나 문장까지 함께 잡혔습니다.")
    add_bullet(doc, "HWP는 텍스트 추출 결과가 0인 파일이 있어, 해당 문서는 파서만으로 값이 나오지 않았습니다.")
    add_bullet(doc, "XLSX는 헤더 위치와 대상 열을 정확히 고정하지 못해, 값이 다른 열에서 읽히는 경우가 생겼습니다.")
    add_bullet(doc, "이미 만들어진 결과 파일에서 불완전한 중간 산출물을 다시 잘라 쓰면서, 오염이 그대로 남았습니다.")

    h = doc.add_paragraph()
    r = h.add_run("3. 증거")
    set_font(r, size=12, bold=True)
    add_bullet(doc, "첫 행은 헤더인데, 2행부터 이미 잘못된 숫자·문구가 섞여 있었습니다.")
    add_bullet(doc, "검증 열(예: 검증상태, 검증메모, 비고)이 붙은 중간 산출물이 그대로 사용되었습니다.")
    add_bullet(doc, "최종 결과 파일은 18개 항목을 위한 깨끗한 재추출본이 아니라, 기존 파일의 잘린 복사본이었습니다.")

    h = doc.add_paragraph()
    r = h.add_run("4. 재발 방지 기준")
    set_font(r, size=12, bold=True)
    add_bullet(doc, "학교 1개, 파일 1개, 원문 1페이지를 먼저 검증한 뒤 전체로 확장합니다.")
    add_bullet(doc, "XLSX는 헤더 행을 먼저 고정하고, 각 항목의 열을 이름으로 매핑합니다.")
    add_bullet(doc, "PDF는 텍스트 추출과 OCR을 분리하고, 키워드가 있는 페이지만 읽습니다.")
    add_bullet(doc, "HWP는 별도 변환 경로를 두고, 텍스트가 비면 즉시 예외 처리합니다.")

    h = doc.add_paragraph()
    r = h.add_run("5. 최종 판단")
    set_font(r, size=12, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(
        "즉, 문제는 데이터가 없어서가 아니라 '추출 규칙, 열 매핑, 소스 식별'이 함께 무너진 것입니다. "
        "이 상태의 파일은 완성본으로 볼 수 없습니다."
    )
    set_font(r)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
