import io
import json
import os
import sqlite3
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
import streamlit as st
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DB_PATH = BASE_DIR / "admission_diag.db"

SUSI_2026 = DATA_DIR / "susi_explorer.csv"
CUTOFF_2026 = DATA_DIR / "admission_cutoffs.csv"
SUSI_2027 = DATA_DIR / "susi_explorer_2027.csv"
CUTOFF_2027 = DATA_DIR / "admission_cutoffs_2027.csv"

GRADE_OPTIONS = ["고1", "고2", "고3", "N수이상"]
ADMISSION_TYPES = ["학생부교과", "학생부종합", "논술", "실기/실적", "기타"]
SEOUL_TOP15 = {
    "서울대", "연세대", "고려대", "서강대", "성균관대",
    "한양대", "중앙대", "경희대", "한국외대", "서울시립대",
    "이화여대", "건국대", "동국대", "홍익대", "숙명여대",
}


@dataclass
class SupportChoice:
    university: str
    department: str
    admission_type: str
    track_name: str
    diag_level: str
    diag_reason: str


def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)


def init_db() -> None:
    ensure_dirs()
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS diagnosis_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                consultant_name TEXT,
                student_name TEXT,
                school_name TEXT,
                grade TEXT,
                student_phone TEXT,
                email TEXT,
                parent_phone TEXT,
                student_index REAL,
                pdf_summary TEXT,
                holistic_score REAL,
                outcome_json TEXT
            )
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS support_choices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id INTEGER NOT NULL,
                support_no INTEGER NOT NULL,
                university TEXT,
                department TEXT,
                admission_type TEXT,
                track_name TEXT,
                diag_level TEXT,
                diag_reason TEXT,
                FOREIGN KEY(session_id) REFERENCES diagnosis_sessions(id)
            )
            """
        )
        cur.execute("CREATE INDEX IF NOT EXISTS idx_sessions_created_at ON diagnosis_sessions(created_at)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_choices_session_id ON support_choices(session_id)")
        conn.commit()
    finally:
        conn.close()


def load_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, encoding="utf-8-sig")


def to_int_year(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(-1).astype(int)


def normalize_university_name(name: str) -> str:
    n = str(name or "").strip().replace(" ", "")
    if "(" in n:
        n = n.split("(", 1)[0]
    n = n.replace("대학교", "대")
    return n


def collect_2027_universities(base_susi: pd.DataFrame, base_cutoff: pd.DataFrame) -> List[str]:
    universities = set()
    for df in [base_susi, base_cutoff]:
        if not df.empty and "year" in df.columns and "university" in df.columns:
            y = to_int_year(df["year"])
            universities.update(df.loc[y == 2027, "university"].dropna().astype(str).tolist())
    for path in [SUSI_2027, CUTOFF_2027]:
        df = load_csv(path)
        if not df.empty and "university" in df.columns:
            universities.update(df["university"].dropna().astype(str).tolist())
    return sorted(normalize_university_name(x) for x in universities if str(x).strip())


def build_2026_frames(base_susi: pd.DataFrame, base_cutoff: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, List[str]]:
    susi_2026 = base_susi.copy()
    cutoff_2026 = base_cutoff.copy()
    if not susi_2026.empty and "year" in susi_2026.columns:
        susi_2026 = susi_2026.loc[to_int_year(susi_2026["year"]) == 2026].copy()
    if not cutoff_2026.empty and "year" in cutoff_2026.columns:
        cutoff_2026 = cutoff_2026.loc[to_int_year(cutoff_2026["year"]) == 2026].copy()

    susi_2027 = load_csv(SUSI_2027)
    cutoff_2027 = load_csv(CUTOFF_2027)
    if not susi_2027.empty and "year" in susi_2027.columns:
        susi_2027 = susi_2027.loc[to_int_year(susi_2027["year"]) == 2027].copy()
    if not cutoff_2027.empty and "year" in cutoff_2027.columns:
        cutoff_2027 = cutoff_2027.loc[to_int_year(cutoff_2027["year"]) == 2027].copy()

    updated_schools = sorted(
        {normalize_university_name(x) for x in susi_2027.get("university", pd.Series(dtype=str)).dropna().astype(str).tolist()}
    )

    susi_df = susi_2026.copy()
    if not susi_2027.empty and updated_schools and "university" in susi_df.columns:
        base_keep = susi_df[~susi_df["university"].astype(str).map(normalize_university_name).isin(updated_schools)].copy()
        susi_2027["year"] = 2026
        susi_df = pd.concat([base_keep, susi_2027], ignore_index=True)

    cutoff_df = cutoff_2026.copy()
    if not cutoff_2027.empty and updated_schools and "university" in cutoff_df.columns:
        base_keep_cut = cutoff_df[~cutoff_df["university"].astype(str).map(normalize_university_name).isin(updated_schools)].copy()
        cutoff_2027["year"] = 2026
        cutoff_df = pd.concat([base_keep_cut, cutoff_2027], ignore_index=True)

    return susi_df, cutoff_df, updated_schools


def extract_pdf_text(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def analyze_holistic_5level(pdf_text: str) -> Tuple[int, Dict[str, int], str]:
    text = (pdf_text or "").strip()
    normalized = text.replace(" ", "").replace("\n", "")
    length_bonus = min(18, len(normalized) // 700) if normalized else 0
    baseline = 45 if normalized else 35
    keywords = {
        "학업역량": ["세부능력", "교과", "성취", "수업", "학습", "심화"],
        "전공적합성": ["진로", "전공", "탐구", "관련", "연계", "프로젝트"],
        "자기주도성": ["주도", "기획", "문제해결", "실행", "탐색", "자기주도"],
        "공동체역량": ["협력", "소통", "리더", "봉사", "배려", "팀"],
        "발전가능성": ["성장", "피드백", "개선", "도전", "변화", "발전"],
    }
    per_item: Dict[str, int] = {}
    for key, words in keywords.items():
        cnt = sum(normalized.count(w) for w in words)
        section_bonus = 8 if key in normalized else 0
        per_item[key] = max(20, min(100, baseline + cnt * 7 + length_bonus + section_bonus))

    avg = int(sum(per_item.values()) / len(per_item)) if per_item else baseline
    if avg >= 85:
        level = 5
    elif avg >= 72:
        level = 4
    elif avg >= 60:
        level = 3
    elif avg >= 48:
        level = 2
    else:
        level = 1

    if not normalized:
        summary = "PDF 텍스트 추출이 제한되어 중립 기준으로 산정했습니다. (스캔본은 OCR 권장)"
    else:
        summary = f"서울 15개 대학 학생부종합 평가축 기준으로 5단계 중 {level}단계로 추정됩니다."
    return level, per_item, summary


def has_openai_key() -> bool:
    key = os.environ.get("OPENAI_API_KEY", "").strip()
    session_key = str(st.session_state.get("openai_api_key", "")).strip() if hasattr(st, "session_state") else ""
    return bool(key or session_key)


def build_ai_holistic_comment(pdf_text: str, level: int, detail: Dict[str, int]) -> str:
    if not has_openai_key():
        return ""
    try:
        from openai import OpenAI
    except Exception:
        return ""

    api_key = str(st.session_state.get("openai_api_key", "")).strip() or os.environ.get("OPENAI_API_KEY", "").strip()
    client = OpenAI(api_key=api_key)
    detail_text = ", ".join([f"{k}:{v}" for k, v in detail.items()])
    prompt = (
        "학생부 분석 보조 코멘트를 간결하게 작성해줘. "
        "출력 형식: 1) 강점 2줄 2) 보완점 2줄 3) 다음 액션 2줄. "
        f"기본분석 단계={level}, 항목점수={detail_text}. "
        f"학생부 텍스트 샘플={pdf_text[:2800]}"
    )
    try:
        res = client.responses.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            input=prompt,
            max_output_tokens=260,
        )
        return (res.output_text or "").strip()
    except Exception:
        return ""


def rating_4level(student_index: float, cutoff50: float, cutoff70: float) -> Tuple[str, str]:
    if student_index >= cutoff50:
        return "상향 가능", f"학생 지표({student_index:.1f})가 50% 컷({cutoff50:.1f}) 이상입니다."
    if student_index >= cutoff70:
        return "적정", f"학생 지표({student_index:.1f})가 70% 컷({cutoff70:.1f}) 이상입니다."
    if student_index >= cutoff70 - 5:
        return "도전", f"학생 지표({student_index:.1f})가 70% 컷({cutoff70:.1f})에 근접합니다."
    return "하향 권장", f"학생 지표({student_index:.1f})가 70% 컷({cutoff70:.1f}) 대비 낮습니다."


def criteria_basis_for_university(university: str) -> str:
    uni = normalize_university_name(university)
    if uni in {normalize_university_name(x) for x in SEOUL_TOP15}:
        return f"{university} 기준"
    return "경희대 기준(대체)"


def get_cutoff(cutoffs: pd.DataFrame, uni: str, dept: str, atype: str, track: str) -> Tuple[float, float]:
    if cutoffs.empty:
        return 85.0, 80.0
    uni_norm = normalize_university_name(uni)
    uni_series = cutoffs["university"].astype(str).map(normalize_university_name)
    hit = cutoffs[
        (uni_series == uni_norm)
        & (cutoffs["department"].astype(str) == str(dept))
        & (cutoffs["admission_type"].astype(str) == str(atype))
        & (cutoffs["track_name"].astype(str) == str(track))
    ]
    if hit.empty:
        return 85.0, 80.0
    c50 = hit.loc[hit["percentile_type"] == 50, "cutoff_score"]
    c70 = hit.loc[hit["percentile_type"] == 70, "cutoff_score"]
    return float(c50.iloc[0] if not c50.empty else 85.0), float(c70.iloc[0] if not c70.empty else 80.0)


def save_session(payload: Dict, choices: List[SupportChoice]) -> int:
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.cursor()
        payload_out = dict(payload)
        payload_out["supports"] = [c.__dict__ for c in choices]
        cur.execute(
            """
            INSERT INTO diagnosis_sessions (
                created_at, consultant_name, student_name, school_name, grade,
                student_phone, email, parent_phone, student_index, pdf_summary,
                holistic_score, outcome_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                datetime.now().isoformat(timespec="seconds"),
                payload.get("consultant_name", ""),
                payload.get("student_name", ""),
                payload.get("school_name", ""),
                payload.get("grade", ""),
                payload.get("student_phone", ""),
                payload.get("email", ""),
                payload.get("parent_phone", ""),
                float(payload.get("student_index", 0)),
                payload.get("pdf_summary", ""),
                float(payload.get("holistic_score", 0)),
                json.dumps(payload_out, ensure_ascii=False),
            ),
        )
        sid = int(cur.lastrowid)
        for i, row in enumerate(choices, start=1):
            cur.execute(
                """
                INSERT INTO support_choices (
                    session_id, support_no, university, department, admission_type,
                    track_name, diag_level, diag_reason
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    sid,
                    i,
                    row.university,
                    row.department,
                    row.admission_type,
                    row.track_name,
                    row.diag_level,
                    row.diag_reason,
                ),
            )
        conn.commit()
        return sid
    finally:
        conn.close()


def report_markdown(payload: Dict, choices: List[SupportChoice], holistic_detail: Dict[str, int]) -> str:
    lines = [
        "# 나의 입시 위치 진단 보고서",
        f"- 생성시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 1) 개인정보",
        f"- 컨설턴트: {payload['consultant_name']}",
        f"- 학생명: {payload['student_name']}",
        f"- 학교명: {payload['school_name']}",
        f"- 학년: {payload['grade']}",
        f"- 학생 연락처: {payload['student_phone']}",
        f"- 이메일: {payload['email']}",
        f"- 학부모 연락처: {payload['parent_phone']}",
        "",
        "## 2) 학생부 분석",
        f"- 종합 단계(5단계): {payload['holistic_level']}",
    ]
    for k, v in holistic_detail.items():
        lines.append(f"- {k}: {v}점")
    if payload.get("ai_comment", "").strip():
        lines.extend(["", "## 2-1) AI 보조 분석 코멘트", payload["ai_comment"]])
    lines.extend(
        [
            "",
            "## 2-2) 평가 항목 상세",
            "- 학업역량: 교과 성취, 수업 태도, 심화 학습",
            "- 전공적합성: 전공 연계 선택/활동, 탐구 맥락",
            "- 자기주도성: 기획-실행-개선 흐름",
            "- 공동체역량: 협업/소통/리더십/배려",
            "- 발전가능성: 학년별 성장, 피드백 반영",
            "",
            "## 3) 지원희망대학 평가(최대 6개, 4단계)",
        ]
    )
    for i, c in enumerate(choices, start=1):
        lines.append(f"- {i}. {c.university} / {c.department} / {c.admission_type} / {c.track_name}")
        lines.append(f"  - 기준: {criteria_basis_for_university(c.university)}")
        lines.append(f"  - 결과: {c.diag_level}")
        lines.append(f"  - 근거: {c.diag_reason}")
    lines.extend(
        [
            "",
            "## 4) 종합 코멘트",
            "- 서울 15개 대학 학생부종합 평가축으로 1차 정량화했습니다.",
            "- 서울 15개 대학 외 지원대학은 경희대 평가축을 대체 적용했습니다.",
            "- 2027 검색기에 존재하는 대학은 2026 진단대상에서 제외했습니다.",
        ]
    )
    return "\n".join(lines)


def markdown_to_pdf_bytes(text: str) -> bytes:
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], fontName="HYSMyeongJo-Medium", fontSize=18)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="HYSMyeongJo-Medium", fontSize=10, leading=14)
    story = []
    for line in text.splitlines():
        if line.startswith("# "):
            story.append(Paragraph(line[2:], title))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe if safe else " ", body))
        story.append(Spacer(1, 4))
    doc.build(story)
    buf.seek(0)
    return buf.read()


def markdown_to_docx_bytes(text: str) -> bytes:
    try:
        from docx import Document
    except Exception:
        return text.encode("utf-8")
    doc = Document()
    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def apply_style(ui_scale: float, page_max_width: int) -> None:
    css = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;500;700;800&display=swap');

:root {
  --ui-scale: __UI_SCALE__;
  --page-max: __PAGE_MAX__px;
}

* {font-family: "Noto Sans KR", sans-serif;}
section.main > div.block-container {max-width: var(--page-max); padding-top: calc(14px * var(--ui-scale)); padding-bottom: calc(28px * var(--ui-scale));}
.stApp {background: radial-gradient(1200px 520px at 20% -20%, #12316b 0%, #0a1634 36%, #050d1f 70%, #030913 100%);}
h1, h2, h3, label, p, div, span {letter-spacing: -0.15px;}

.top-meta {font-size: calc(18px * var(--ui-scale)); font-weight: 700; color: #f8fafc; margin-bottom: 2px;}
.top-sub {font-size: calc(14px * var(--ui-scale)); color: #94a3b8; margin-bottom: calc(14px * var(--ui-scale));}

.step-grid {display: grid; grid-template-columns: repeat(4, 1fr); gap: calc(16px * var(--ui-scale)); margin: calc(8px * var(--ui-scale)) 0 calc(18px * var(--ui-scale));}
.step-item {
  height: calc(58px * var(--ui-scale)); border-radius: calc(14px * var(--ui-scale)); display:flex; align-items:center; justify-content:center;
  font-size: calc(27px * var(--ui-scale)); font-weight: 700; letter-spacing: -0.2px;
  color: #8fa1bc; background: #e5e7eb; border: 1px solid #cbd5e1;
}
.step-item.active {
  color: #ffffff; background: linear-gradient(90deg,#1d4ed8 0%, #2563eb 100%);
  border: 1px solid #2563eb; box-shadow: 0 8px 22px rgba(37,99,235,0.32);
}

.rule-line {height:1px; width:100%; background: rgba(148,163,184,0.35); margin: calc(8px * var(--ui-scale)) 0 calc(20px * var(--ui-scale));}

.section-card {
  border: 1px solid rgba(148,163,184,0.35); border-radius: calc(14px * var(--ui-scale)); padding: calc(18px * var(--ui-scale)) calc(18px * var(--ui-scale)) calc(20px * var(--ui-scale));
  background: linear-gradient(180deg, rgba(9,18,38,0.82), rgba(3,9,19,0.86));
}
.section-banner {
  height: calc(46px * var(--ui-scale)); border-left: 5px solid #3b82f6; border-radius: 0 calc(8px * var(--ui-scale)) calc(8px * var(--ui-scale)) 0;
  background: linear-gradient(90deg, rgba(226,232,240,0.95), rgba(100,116,139,0.32), rgba(30,41,59,0.04));
  color: #0f172a; display:flex; align-items:center; padding: 0 calc(14px * var(--ui-scale));
  font-size: calc(30px * var(--ui-scale)); font-weight: 800; letter-spacing: -0.35px; margin-bottom: calc(14px * var(--ui-scale));
}

[data-testid="stTextInputRootElement"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] div[data-baseweb="select"] > div {
  min-height: calc(52px * var(--ui-scale)) !important; border-radius: calc(12px * var(--ui-scale)) !important; font-size: calc(27px * var(--ui-scale)) !important;
  background: rgba(15,23,42,0.88) !important; color: #f8fafc !important;
  border: 1px solid rgba(203,213,225,0.6) !important;
}
[data-testid="stTextArea"] textarea {min-height: calc(116px * var(--ui-scale)) !important;}
[data-testid="stFileUploaderDropzone"] {
  border-radius: calc(14px * var(--ui-scale)) !important; background: rgba(30,41,59,0.6) !important;
  border: 1px dashed rgba(148,163,184,0.52) !important;
}

[data-testid="stButton"] button,
[data-testid="stDownloadButton"] button {
  min-height: calc(52px * var(--ui-scale)) !important; border-radius: calc(14px * var(--ui-scale)) !important; font-size: calc(28px * var(--ui-scale)) !important;
  font-weight: 700 !important; letter-spacing: -0.2px;
  border: 1px solid rgba(37,99,235,0.7) !important; color: #ffffff !important;
  background: linear-gradient(90deg, #1d4ed8 0%, #2563eb 100%) !important;
  box-shadow: 0 6px 16px rgba(37,99,235,0.25);
}
[data-testid="stButton"] button:hover,
[data-testid="stDownloadButton"] button:hover {
  filter: brightness(1.05);
}

.metric-row {margin-bottom: 10px;}
.metric-card {border:1px solid rgba(148,163,184,0.25); border-radius: calc(12px * var(--ui-scale)); padding: calc(12px * var(--ui-scale)) calc(14px * var(--ui-scale)); background: rgba(2,6,23,0.58);}
.metric-title {font-size:calc(24px * var(--ui-scale)); color:#94a3b8; font-weight:500;}
.metric-value {font-size:calc(34px * var(--ui-scale)); font-weight:800; color:#e2e8f0; margin-top:2px;}

@media (max-width: 1100px) {
  .step-grid {grid-template-columns: repeat(2, 1fr);}
  .section-banner {font-size: calc(24px * var(--ui-scale));}
}
</style>
    """
    css = css.replace("__UI_SCALE__", f"{ui_scale:.2f}").replace("__PAGE_MAX__", str(page_max_width))
    st.markdown(
        css,
        unsafe_allow_html=True,
    )


def render_stepper(step: int) -> None:
    names = ["1. 개인정보", "2. 학생부 분석", "3. 지원 대학", "4. 종합 보고서"]
    html = []
    for idx, name in enumerate(names, start=1):
        cls = "step-item active" if idx == step else "step-item"
        html.append(f"<div class='{cls}'>{name}</div>")
    st.markdown("<div class='step-grid'>" + "".join(html) + "</div>", unsafe_allow_html=True)


def section_banner(text: str) -> None:
    st.markdown(f"<div class='section-banner'>📋 {text}</div>", unsafe_allow_html=True)


def init_state() -> None:
    ss = st.session_state
    ss.setdefault("step", 1)
    ss.setdefault("profile", {})
    ss.setdefault("holistic", {})
    ss.setdefault("choices", [])
    ss.setdefault("saved_session_id", None)
    ss.setdefault("ui_scale_mode", "표준 (100%)")
    ss.setdefault("ui_width_mode", "1400px")


def main() -> None:
    st.set_page_config(page_title="나의 입시 위치 진단 서비스", layout="wide")
    init_db()
    init_state()

    base_susi = load_csv(SUSI_2026)
    base_cutoff = load_csv(CUTOFF_2026)
    susi_df, cutoff_df, updated_2027 = build_2026_frames(base_susi, base_cutoff)

    st.markdown("<div class='top-meta'>나의 입시 위치 진단 서비스</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='top-sub'>단계형 진단: 개인정보 → 학생부 분석 → 지원희망대학 평가(최대 6개) → 종합결과/보고서</div>",
        unsafe_allow_html=True,
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(
            f"<div class='metric-card'><div class='metric-title'>2026 지원 데이터</div><div class='metric-value'>{len(susi_df):,}건</div></div>",
            unsafe_allow_html=True,
        )
    with c2:
        st.markdown(
            f"<div class='metric-card'><div class='metric-title'>2026 컷 데이터</div><div class='metric-value'>{len(cutoff_df):,}건</div></div>",
            unsafe_allow_html=True,
        )
    with c3:
        st.markdown(
            f"<div class='metric-card'><div class='metric-title'>2027 업데이트 대학</div><div class='metric-value'>{len(updated_2027):,}개</div></div>",
            unsafe_allow_html=True,
        )

    with st.sidebar:
        st.subheader("데이터 상태")
        st.write(f"2026 지원 데이터: {len(susi_df)}건")
        st.write(f"2026 컷 데이터: {len(cutoff_df)}건")
        st.write(f"2027 업데이트 대학: {len(updated_2027)}개")
        st.caption(f"DB: {DB_PATH}")
        st.markdown("---")
        st.subheader("API 설정")
        st.text_input("OpenAI API Key (선택)", type="password", key="openai_api_key")
        if st.session_state.get("openai_api_key", "").strip():
            os.environ["OPENAI_API_KEY"] = st.session_state.openai_api_key.strip()
        st.caption("API 상태: " + ("연결됨" if has_openai_key() else "미설정 (규칙 기반 분석)"))
        st.markdown("---")
        st.subheader("화면 설정")
        st.selectbox("화면 배율", ["작게 (90%)", "표준 (100%)", "크게 (115%)"], key="ui_scale_mode")
        st.selectbox("최대 폭", ["1200px", "1400px", "1600px"], key="ui_width_mode")

        if updated_2027:
            with st.expander("2027 기준으로 업데이트된 대학 목록", expanded=False):
                st.write(", ".join(updated_2027[:120]))
                if len(updated_2027) > 120:
                    st.caption(f"... 외 {len(updated_2027)-120}개")

        admin_mode = st.checkbox("관리자 모드")
        if admin_mode:
            st.markdown("---")
            up1 = st.file_uploader("2026 수시탐색기 CSV", type=["csv"], key="up_susi_2026")
            if up1:
                SUSI_2026.write_bytes(up1.getvalue())
                st.success("2026 수시탐색기 반영 완료")
            up2 = st.file_uploader("2026 컷 CSV", type=["csv"], key="up_cut_2026")
            if up2:
                CUTOFF_2026.write_bytes(up2.getvalue())
                st.success("2026 컷 반영 완료")
            up3 = st.file_uploader("2027 수시탐색기 CSV", type=["csv"], key="up_susi_2027")
            if up3:
                SUSI_2027.write_bytes(up3.getvalue())
                st.success("2027 수시탐색기 반영 완료")
            up4 = st.file_uploader("2027 컷 CSV(선택)", type=["csv"], key="up_cut_2027")
            if up4:
                CUTOFF_2027.write_bytes(up4.getvalue())
                st.success("2027 컷 반영 완료")

    scale_map = {"작게 (90%)": 0.90, "표준 (100%)": 1.00, "크게 (115%)": 1.15}
    width_map = {"1200px": 1200, "1400px": 1400, "1600px": 1600}
    ui_scale = scale_map.get(st.session_state.ui_scale_mode, 1.00)
    ui_width = width_map.get(st.session_state.ui_width_mode, 1400)
    apply_style(ui_scale=ui_scale, page_max_width=ui_width)

    render_stepper(st.session_state.step)
    st.markdown("<div class='rule-line'></div>", unsafe_allow_html=True)
    st.markdown("<div class='section-card'>", unsafe_allow_html=True)

    if st.session_state.step == 1:
        section_banner("1단계 - 개인정보 입력")
        a, b, c = st.columns(3)
        with a:
            student_name = st.text_input("학생 이름 *", value=st.session_state.profile.get("student_name", ""))
            student_phone = st.text_input("학생 연락처", value=st.session_state.profile.get("student_phone", ""))
            parent_phone = st.text_input("학부모 연락처", value=st.session_state.profile.get("parent_phone", ""))
        with b:
            school_name = st.text_input("학교명 *", value=st.session_state.profile.get("school_name", ""))
            default_grade = st.session_state.profile.get("grade", "고3")
            grade = st.selectbox("현재 학년", GRADE_OPTIONS, index=GRADE_OPTIONS.index(default_grade) if default_grade in GRADE_OPTIONS else 2)
            student_index = st.slider("학생 지표 (50~100)", 50.0, 100.0, float(st.session_state.profile.get("student_index", 82.0)), 0.5)
        with c:
            email = st.text_input("이메일 * (보고서 발송용)", value=st.session_state.profile.get("email", ""))
            consultant_name = st.text_input("담당 컨설턴트", value=st.session_state.profile.get("consultant_name", ""))
            memo = st.text_area("메모", value=st.session_state.profile.get("memo", ""), height=120)

        c_prev, c_next = st.columns([1, 1])
        with c_next:
            if st.button("다음 단계 →", use_container_width=True):
                if not student_name.strip() or not school_name.strip() or not email.strip():
                    st.error("필수 항목(학생 이름, 학교명, 이메일)을 입력해 주세요.")
                else:
                    st.session_state.profile = {
                        "consultant_name": consultant_name.strip(),
                        "student_name": student_name.strip(),
                        "school_name": school_name.strip(),
                        "grade": grade,
                        "student_phone": student_phone.strip(),
                        "email": email.strip(),
                        "parent_phone": parent_phone.strip(),
                        "student_index": float(student_index),
                        "memo": memo.strip(),
                    }
                    st.session_state.step = 2
                    st.rerun()

    elif st.session_state.step == 2:
        section_banner("2단계 - 학생부 분석")
        uploaded = st.file_uploader("학생부 PDF 업로드", type=["pdf"])
        ai_enabled = st.toggle("AI 보조 분석 사용 (OPENAI_API_KEY 필요)", value=st.session_state.get("ai_enabled", False))
        st.session_state.ai_enabled = ai_enabled
        if ai_enabled and not has_openai_key():
            st.warning("OPENAI_API_KEY가 없어 규칙 기반 분석만 수행됩니다.")
        x, y = st.columns(2)
        with x:
            if st.button("학생부 분석 실행", use_container_width=True):
                pdf_text = extract_pdf_text(uploaded) if uploaded else ""
                level, detail, summary = analyze_holistic_5level(pdf_text)
                ai_comment = ""
                if ai_enabled and has_openai_key():
                    ai_comment = build_ai_holistic_comment(pdf_text, level, detail)
                st.session_state.holistic = {
                    "level": level,
                    "detail": detail,
                    "summary": summary,
                    "ai_enabled": ai_enabled,
                    "ai_comment": ai_comment,
                }
                st.success(summary)
        with y:
            if st.button("다음 단계 →", use_container_width=True):
                if not st.session_state.holistic:
                    st.error("먼저 학생부 분석을 실행해 주세요.")
                else:
                    st.session_state.step = 3
                    st.rerun()

        if st.session_state.holistic:
            st.info(st.session_state.holistic["summary"])
            for k, v in st.session_state.holistic["detail"].items():
                st.write(f"- {k}: {v}점")
            if st.session_state.holistic.get("ai_comment", "").strip():
                st.markdown("**AI 보조 코멘트**")
                st.write(st.session_state.holistic["ai_comment"])
        if st.button("← 이전 단계"):
            st.session_state.step = 1
            st.rerun()

    elif st.session_state.step == 3:
        section_banner("3단계 - 지원희망대학 평가 (최대 6개)")
        if susi_df.empty:
            st.error("2026 수시 데이터가 비어 있습니다. data/susi_explorer.csv를 확인해 주세요.")
        else:
            uni_options = sorted(susi_df["university"].dropna().astype(str).unique().tolist())

            def _support_filled(idx: int) -> bool:
                u, d, a, t = st.session_state.get(f"support_{idx}", ("", "", "", ""))
                return bool(u and d and a and t)

            for i in range(1, 7):
                unlocked = True
                if i == 3:
                    unlocked = _support_filled(1) and _support_filled(2)
                elif i >= 4:
                    unlocked = _support_filled(i - 1)
                default_open = i in (1, 2)
                if i == 3 and unlocked:
                    default_open = True

                with st.expander(f"지원 {i}", expanded=default_open):
                    if not unlocked:
                        st.info(f"지원 {i-1} 입력 완료 후 열립니다.")
                        continue

                    c1, c2, c3, c4 = st.columns(4)
                    with c1:
                        uni_sel = st.selectbox(f"학교 {i}", ["직접입력"] + uni_options, key=f"uni_{i}_sel")
                        uni_val = st.text_input(f"학교 직접입력 {i}", key=f"uni_{i}_txt") if uni_sel == "직접입력" else uni_sel

                    filtered = susi_df[susi_df["university"] == uni_val] if uni_val else susi_df
                    dept_options = sorted(filtered["department"].dropna().astype(str).unique().tolist()) if not filtered.empty else []
                    track_options = sorted(filtered["track_name"].dropna().astype(str).unique().tolist()) if not filtered.empty else []

                    with c2:
                        dept_sel = st.selectbox(f"학과 {i}", ["직접입력"] + dept_options, key=f"dept_{i}_sel")
                        dept_val = st.text_input(f"학과 직접입력 {i}", key=f"dept_{i}_txt") if dept_sel == "직접입력" else dept_sel
                    with c3:
                        atype = st.selectbox(f"전형유형 {i}", ADMISSION_TYPES, key=f"atype_{i}")
                    with c4:
                        track_sel = st.selectbox(f"전형명 {i}", ["직접입력"] + track_options, key=f"track_{i}_sel")
                        track_val = st.text_input(f"전형명 직접입력 {i}", key=f"track_{i}_txt") if track_sel == "직접입력" else track_sel

                    st.session_state[f"support_{i}"] = (uni_val.strip(), dept_val.strip(), atype.strip(), track_val.strip())

            p1, p2, p3 = st.columns([1, 1, 1])
            with p2:
                if st.button("평가 실행 후 결과 보기 →", use_container_width=True):
                    choices: List[SupportChoice] = []
                    for i in range(1, 7):
                        uni, dept, atype, track = st.session_state.get(f"support_{i}", ("", "", "", ""))
                        if not (uni and dept and atype and track):
                            continue
                        c50, c70 = get_cutoff(cutoff_df, uni, dept, atype, track)
                        level, reason = rating_4level(float(st.session_state.profile["student_index"]), c50, c70)
                        basis = criteria_basis_for_university(uni)
                        reason_with_basis = f"{reason} | 기준: {basis}"
                        choices.append(SupportChoice(uni, dept, atype, track, level, reason_with_basis))
                    if not choices:
                        st.error("최소 1개 지원 항목을 입력해 주세요.")
                    else:
                        st.session_state.choices = choices
                        st.session_state.saved_session_id = None
                        st.session_state.step = 4
                        st.rerun()
            with p1:
                if st.button("← 이전 단계", use_container_width=True):
                    st.session_state.step = 2
                    st.rerun()

    else:
        section_banner("4단계 - 종합결과 보기 및 보고서 다운로드")
        if not st.session_state.choices:
            st.warning("먼저 3단계 평가를 완료해 주세요.")
        else:
            profile = st.session_state.profile
            holistic = st.session_state.holistic
            payload = {
                "consultant_name": profile["consultant_name"],
                "student_name": profile["student_name"],
                "school_name": profile["school_name"],
                "grade": profile["grade"],
                "student_phone": profile["student_phone"],
                "email": profile["email"],
                "parent_phone": profile["parent_phone"],
                "student_index": profile["student_index"],
                "pdf_summary": holistic.get("summary", ""),
                "holistic_score": sum(holistic.get("detail", {}).values()) / max(1, len(holistic.get("detail", {}))),
                "holistic_level": holistic.get("level", 0),
                "ai_comment": holistic.get("ai_comment", ""),
            }

            st.markdown("**학생부 분석 결과**")
            st.info(holistic.get("summary", "학생부 분석 결과가 없습니다."))

            st.markdown("**지원희망대학 평가 결과**")
            for i, row in enumerate(st.session_state.choices, start=1):
                st.write(f"{i}. {row.university} / {row.department} / {row.admission_type} / {row.track_name}")
                st.caption(f"평가: {row.diag_level} | 근거: {row.diag_reason}")

            md = report_markdown(payload, st.session_state.choices, holistic.get("detail", {}))
            pdf_bytes = markdown_to_pdf_bytes(md)
            docx_bytes = markdown_to_docx_bytes(md)
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            with st.expander("종합보고서 form 미리보기", expanded=True):
                st.text_area(
                    "아래 내용이 PDF/DOCX로 저장됩니다.",
                    value=md,
                    height=360,
                    key="final_report_preview",
                )

            d1, d2, d3 = st.columns(3)
            with d1:
                st.download_button(
                    "보고서 다운로드 (PDF)",
                    data=pdf_bytes,
                    file_name=f"admission_report_{stamp}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            with d2:
                st.download_button(
                    "보고서 다운로드 (DOCX)",
                    data=docx_bytes,
                    file_name=f"admission_report_{stamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with d3:
                if st.button("DB 저장", use_container_width=True):
                    sid = save_session(payload, st.session_state.choices)
                    st.session_state.saved_session_id = sid
                    st.success(f"DB 저장 완료 (session_id={sid})")

            if st.session_state.saved_session_id:
                st.caption(f"최근 저장 session_id: {st.session_state.saved_session_id}")

            if st.button("← 이전 단계"):
                st.session_state.step = 3
                st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
